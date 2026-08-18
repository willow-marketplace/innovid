# /// script
# requires-python = ">=3.9"
# dependencies = [
#   "openpyxl>=3.1.0",
#   "pymupdf>=1.24.0",
#   "python-docx>=1.1.0",
# ]
# ///
"""Turn an uploaded spreadsheet or document into carta-issuance's `knowns.rows`.

The file feeds the FRONT of the existing pipeline: this script's whole output
contract is the same `knowns.rows` shape build_config.py already consumes, so
the config panel opens prefilled and every downstream gate (Phase 1 resolve →
Phase 1.5 save+validate → Phase 2 review → Phase 3 mutate) runs untouched.
Nothing here writes to Carta.

Two modes, picked by extension:

  .xlsx .xlsm .csv .tsv   spreadsheet — fully deterministic. Detects the header
                          row (Carta's importer template puts instruction prose
                          in row 1), maps headers via COLUMN_SYNONYMS, coerces
                          values, and resolves names to ids against the
                          reference payload.
  .pdf .docx              document — extracts text to _import_text.txt and
                          STOPS. Prose has no fixed layout, so building rows
                          from it is the model's job; doing it here would mean
                          guessing silently. The model writes rows in this same
                          schema and marks each field confidence "low".

Two rules the rest of the file exists to uphold:

1. **Unresolved is blank, never guessed.** A vesting-schedule or share-class
   cell that doesn't match exactly (case-insensitively) leaves the field unset
   and records an import_note. There is deliberately no fuzzy matching: an
   almost-match issues genuinely wrong terms, and unlike a bad quantity the
   server cannot catch it.
2. **Nothing is dropped silently.** Unmapped columns, skipped rows, and
   uncoercible cells all land in _import_report.json and, per row, in
   `import_notes` so the panel can flag them. A dropped Exercise Price column
   is a wrong-priced grant the user has no way to notice.

Usage:
  uv run parse_upload.py --file <path> [--sheet <name>] \
      [--reference <ref.json>] --out-dir <OUT_DIR>

`--reference` is the same JSON build_config.py takes as `--data` (raw MCP
section envelopes plus `stakeholders`). Omit it to parse without resolving.

Writes to --out-dir:
  _import_knowns.json  {security_type, rows, equity_plan_id?, batch_errors}
  _import_report.json  what was read, what wasn't mapped, what was skipped
  _import_text.txt     document mode only — extracted text for the model

Exit 0 — parsed. Exit 2 — nothing usable (missing file, no header row, no data
rows, ambiguous sheet); the message on stderr says which.
"""
from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

# ── Enums, copied from build_config.py so a parsed value is panel-legal ──

RELATIONSHIP_CHOICES = [
    "Advisor", "Ex-Advisor", "Board member", "Ex-Board member",
    "Consultant", "Ex-Consultant", "Employee", "Ex-Employee", "Executive",
    "Founder", "International Employee", "Ex-International Employee",
    "Investor", "Officer", "Other",
]

OPTION_TYPES = [
    "ISO", "NSO", "INTL", "EMI", "CSOP", "Unapproved",
    "Startup Concessions", "Non-Concessional", "ZEPO",
]

# Security types an importer workbook can carry that this skill does not issue
# (SKILL.md Out of scope). Rows naming one are skipped and reported — never
# coerced into a grant of a different type.
OUT_OF_SCOPE_TYPES = {
    "RSU": "RSUs", "SAR": "SARs", "CBU": "CBUs", "WARRANT": "warrants",
    "RSA": "RSAs", "CONVERTIBLE NOTE": "convertible notes", "SAFE": "SAFEs",
}

GRANT_REASON_CHOICES = [
    "New Hire", "Merit", "Promotion", "Refresh", "Corporate transaction",
    "Relationship change", "Retention", "Advisor", "Consultant", "Board",
    "Performance bonus", "Boxcar grant",
]

# Every key build_config.py reads off a row. A row carrying anything else would
# survive into the save_drafts payload and be rejected server-side.
ROW_KEYS = {
    # shared
    "name", "email", "relationship", "stakeholder_kind", "quantity", "notes",
    "issue_date", "board_approval", "board_approval_date",
    "vesting_template_id", "vesting_start_date", "acceleration_template",
    # option grant
    "option_type", "exercise_price", "document_set_id", "custom_label",
    "grant_reason", "early_exercise", "auto_exercise_at_vest",
    "is_flexible_issue_date", "is_hmrc_notified", "hmrc_notified",
    "is_ato_notified", "grant_expiration_date",
    # certificate
    "share_class_prefix", "price_per_share", "legend_id", "prefix_number",
    "rule_144_mode", "rule_144_date", "rule_144_reason", "cash_paid",
    "debt_canceled", "returned_invested_capital",
}

# ── Header vocabulary ──
#
# Anchored on carta-web's own importer template (example_importer_template_v2_14
# — sheets "Common Certificates" / "Preferred Certificates" / "Equity Plan
# Awards") plus the drafts-v2 in-app grid labels, since those are the two things
# a Carta admin actually has on hand. Keys are normalized (see _norm).
# references/column-map.md documents this table for humans; keep them in step.

COLUMN_SYNONYMS: Dict[str, str] = {}


def _syn(field: str, *headers: str) -> None:
    for h in headers:
        COLUMN_SYNONYMS[_norm(h)] = field


def _norm(s: Any) -> str:
    """Fold a header to a match key: lowercase, punctuation to spaces, collapsed.

    So "Vesting Commencement Date", "vesting_commencement_date" and
    "Vesting  Commencement  Date" are one key.
    """
    return re.sub(r"[^a-z0-9]+", " ", str(s or "").lower()).strip()


# shared
_syn("name", "Name", "Shareholder", "Stakeholder", "Holder", "Grantee",
     "Recipient", "Full Name", "Legal Name", "Employee Name", "Stakeholder Name")
_syn("email", "Email", "Email Address", "Stakeholder Email", "Work Email")
_syn("relationship", "Relationship", "Relationship To Company",
     "Issue Date Relationship")
_syn("stakeholder_kind", "Holder Type", "Stakeholder Type", "Entity Type",
     "Individual Or Non Individual")
_syn("quantity", "Quantity", "Shares", "Number Of Shares", "Share Quantity",
     "Options", "Number Of Options", "Units", "Amount")
_syn("notes", "Notes", "Note", "Comment", "Comments")
_syn("board_approval_date", "Board Approval Date", "Board Approved",
     "Board Approval", "Board Consent Date")
_syn("vesting_template_id", "Vesting Schedule", "Vesting", "Vesting Template")
_syn("vesting_start_date", "Vesting Commencement Date", "Vesting Start Date",
     "Vesting Start", "Vesting Commencement")
_syn("acceleration_template", "Acceleration", "Acceleration Terms",
     "Acceleration Template")

# option grant
_syn("issue_date", "Grant Date", "Issue Date", "Date Issued", "Date Of Grant")
_syn("option_type", "Type", "Award Type", "Option Type", "Grant Type")
_syn("exercise_price", "Exercise Price", "Strike Price", "Exercise Price Per Share")
_syn("document_set_id", "Document Set", "Documents", "Document Template")
_syn("custom_label", "Award ID", "Label", "Custom Label", "Grant Label",
     "Award Label", "Grant Number")
_syn("grant_reason", "Grant Reason", "Reason")
_syn("early_exercise", "Early Exercise", "Allow Early Exercise")
_syn("auto_exercise_at_vest", "Auto Exercise At Vest", "Auto Exercise")
_syn("is_hmrc_notified", "HMRC Notified", "HMRC Notification")
_syn("hmrc_notified", "HMRC Notified Date", "HMRC Notification Date")
_syn("is_ato_notified", "ATO Notified", "ATO Notification")
# NB: the importer template also has an "Exercise Expiration Date" column — the
# post-termination exercise window, a different concept. It stays unmapped (and
# so gets reported), and _norm keeps the two header keys distinct.
_syn("grant_expiration_date", "Expiration Date", "Grant Expiration Date",
     "Expiry Date", "Expiration", "Grant Expiry")
_syn("_equity_plan", "Equity Plan Name", "Equity Plan", "Plan", "Plan Name",
     "Option Plan", "Equity Plan For RSAs")

# certificate
_syn("share_class_prefix", "Share Class", "Class", "Security Class",
     "Share Class Name")
_syn("price_per_share", "Price Per Share", "Purchase Price", "Price",
     "Price Paid Per Share")
_syn("legend_id", "Legend", "Legend Code", "Build Legend")
_syn("prefix_number", "Certificate ID", "Certificate Number", "Cert ID",
     "Certificate No")
_syn("rule_144_date", "Rule 144 Date", "144 Date")
_syn("cash_paid", "Cash Paid", "Total Cash Paid")
_syn("debt_canceled", "Debt Canceled", "Debt Cancelled")
_syn("returned_invested_capital", "Returned Invested Capital")

# Headers that identify the flow but carry no payload field of their own.
_syn("_currency", "Currency")

# Header signatures that decide security_type. Weighted: a header that only ever
# appears on one kind of sheet counts, generic ones (Quantity, Email) don't.
GRANT_SIGNALS = {"option_type", "exercise_price", "document_set_id",
                 "_equity_plan", "is_hmrc_notified", "is_ato_notified"}
CERT_SIGNALS = {"share_class_prefix", "legend_id", "price_per_share",
                "rule_144_date", "prefix_number", "cash_paid", "debt_canceled",
                "returned_invested_capital"}

SHEET_NAME_HINTS = [
    ("equity plan awards", "option_grant"),
    ("awards", "option_grant"),
    ("options", "option_grant"),
    ("grants", "option_grant"),
    ("common certificates", "certificate"),
    ("preferred certificates", "certificate"),
    ("certificates", "certificate"),
]

SPREADSHEET_SUFFIXES = {".xlsx", ".xlsm", ".xltx", ".xltm", ".csv", ".tsv"}
DOCUMENT_SUFFIXES = {".pdf", ".docx"}

# Enough of a header row to be a header row rather than data.
MIN_HEADER_MATCHES = 3
HEADER_SCAN_ROWS = 6


class ParseError(RuntimeError):
    """Nothing usable in the input."""


class AmbiguousInput(RuntimeError):
    """Several candidate sheets; the caller must pick one."""

    def __init__(self, message: str, candidates: Sequence[str]) -> None:
        super().__init__(message)
        self.candidates = list(candidates)


class Sheet:
    """A rectangular grid of raw cell values, plus where it came from."""

    def __init__(self, name: str, rows: List[List[Any]]) -> None:
        self.name = name
        self.rows = rows


class Parsed:
    def __init__(
        self,
        security_type: str,
        rows: List[Dict[str, Any]],
        *,
        source_file: str,
        sheet: Optional[str],
        unmapped_columns: Optional[List[str]] = None,
        skipped_rows: Optional[List[Dict[str, Any]]] = None,
        plan_name: Optional[str] = None,
        equity_plan_id: Optional[str] = None,
        batch_errors: Optional[List[str]] = None,
    ) -> None:
        self.security_type = security_type
        self.rows = rows
        self.source_file = source_file
        self.sheet = sheet
        self.unmapped_columns = unmapped_columns or []
        self.skipped_rows = skipped_rows or []
        self.plan_name = plan_name
        self.equity_plan_id = equity_plan_id
        self.batch_errors = batch_errors or []


# ── Reading ──

def read_sheet(path: Path, sheet: Optional[str] = None) -> Sheet:
    """One sheet as a raw grid. CSV/TSV have exactly one, named for the file."""
    suffix = path.suffix.lower()
    if suffix in {".csv", ".tsv"}:
        delimiter = "\t" if suffix == ".tsv" else ","
        with path.open("r", encoding="utf-8-sig", newline="") as fh:
            rows = [list(r) for r in csv.reader(fh, delimiter=delimiter)]
        return Sheet(path.stem, rows)
    return _read_xlsx_sheets(path, only=sheet)[0]


def _read_xlsx_sheets(path: Path, only: Optional[str] = None) -> List[Sheet]:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        names = wb.sheetnames
        if only is not None:
            match = next((n for n in names if _norm(n) == _norm(only)), None)
            if match is None:
                raise ParseError(
                    "No sheet named {!r}. This file has: {}".format(only, ", ".join(names))
                )
            names = [match]
        out = []
        for name in names:
            ws = wb[name]
            rows = [list(r) for r in ws.iter_rows(values_only=True)]
            out.append(Sheet(name, rows))
        return out
    finally:
        wb.close()


def find_header_row(sheet: Sheet) -> Optional[int]:
    """Index of the header row, or None.

    Carta's importer template puts a paragraph of instructions in row 1 and the
    real headers in row 2, so this scans the first few rows and takes whichever
    maps the most known columns rather than assuming row 0.
    """
    best_idx, best_hits = None, 0
    for idx, row in enumerate(sheet.rows[:HEADER_SCAN_ROWS]):
        hits = len({COLUMN_SYNONYMS[k] for k in (_norm(c) for c in row) if k in COLUMN_SYNONYMS})
        if hits > best_hits:
            best_idx, best_hits = idx, hits
    return best_idx if best_hits >= MIN_HEADER_MATCHES else None


def map_headers(header_row: Sequence[Any]) -> Tuple[Dict[int, str], List[str]]:
    """(column index → field, unmapped header labels).

    A duplicate header keeps the first column: later ones are reported as
    unmapped rather than overwriting a value already read.
    """
    mapping: Dict[int, str] = {}
    unmapped: List[str] = []
    for i, cell in enumerate(header_row):
        label = str(cell).strip() if cell is not None else ""
        if not label:
            continue
        field = COLUMN_SYNONYMS.get(_norm(label))
        if field is None or field in mapping.values():
            unmapped.append(label)
            continue
        mapping[i] = field
    return mapping, unmapped


# ── Value coercion ──

def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.isoformat()[:10]
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


_DATE_FORMATS = ("%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%m-%d-%Y", "%d-%m-%Y",
                 "%Y/%m/%d", "%b %d, %Y", "%d %b %Y", "%B %d, %Y")


def coerce_date(value: Any) -> Optional[str]:
    """→ YYYY-MM-DD, the format every panel date input wants.

    Ambiguous DD/MM vs MM/DD is resolved as MM/DD first (the template's own US
    example) and only re-read as DD/MM when the first field cannot be a month.
    """
    if isinstance(value, (datetime, date)):
        return value.isoformat()[:10]
    raw = _text(value)
    if not raw:
        return None
    raw = raw.split(" 00:00:00")[0].strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(raw, fmt).date().isoformat()
        except ValueError:
            continue
    return None


def coerce_number(value: Any) -> Optional[str]:
    """→ a bare numeric string, or None if it isn't unambiguously a number.

    Tolerates the thousands separators, currency symbols and parenthesised
    negatives that survive a copy-paste out of Excel.

    ``(-100)`` is refused rather than resolved: parentheses already mean negative
    by accounting convention, so the inner minus is either redundant or a second
    negation, and there is no way to tell which. Picking one would put a number
    nobody typed onto a cap table. Returning None hands it to the caller's
    note-and-leave-blank path, which is the same refuse-to-guess rule the name
    resolution follows.
    """
    raw = _text(value)
    if not raw:
        return None
    cleaned = raw.replace(",", "").replace("$", "").replace("£", "").replace("€", "")
    cleaned = cleaned.replace("%", "").strip()
    negative = cleaned.startswith("(") and cleaned.endswith(")")
    if negative:
        cleaned = cleaned[1:-1].strip()
        if cleaned.startswith("-"):
            return None
    if not re.fullmatch(r"-?\d*\.?\d+", cleaned):
        return None
    if negative:
        cleaned = "-" + cleaned
    if cleaned.endswith(".0"):
        cleaned = cleaned[:-2]
    return cleaned


_TRUE = {"yes", "y", "true", "t", "1", "x", "checked"}
_FALSE = {"no", "n", "false", "f", "0", ""}


def coerce_bool(value: Any) -> Optional[bool]:
    raw = _text(value).lower()
    if raw in _TRUE:
        return True
    if raw in _FALSE:
        return False
    return None


def _match_choice(value: Any, choices: Iterable[str]) -> Optional[str]:
    """Exact match on a picklist, case- and punctuation-insensitively. No fuzz."""
    key = _norm(value)
    if not key:
        return None
    for choice in choices:
        if _norm(choice) == key:
            return choice
    return None


def coerce_stakeholder_kind(value: Any) -> Optional[str]:
    key = _norm(value)
    if not key:
        return None
    if key in {"individual", "person", "natural person"}:
        return "INDIVIDUAL"
    if key in {"non individual", "nonindividual", "entity", "organization",
               "organisation", "company", "corporation", "trust", "llc"}:
        return "NON-INDIVIDUAL"
    return None


# ── Row building ──

def _note(row: Dict[str, Any], field: str, raw_value: Any, reason: str) -> None:
    row.setdefault("import_notes", []).append(
        {"field": field, "raw_value": _text(raw_value), "reason": reason}
    )


def _resolved(row: Dict[str, Any], field: str, value: Any) -> None:
    """Set a field AND drop any earlier note against it.

    A note means "the file said something I couldn't turn into a value". Once
    something else supplies that value — the roster, typically — the note is
    stale and must go, because `build_config.py` keys its force-blank behaviour
    off `import_notes` alone. Leaving it would blank a correctly-resolved field
    and tell the admin to pick one, inviting a wrong Individual/Non-individual
    or relationship choice over the cap-table-correct value.
    """
    row[field] = value
    notes = row.get("import_notes")
    if notes:
        kept = [n for n in notes if not (isinstance(n, dict) and n.get("field") == field)]
        if kept:
            row["import_notes"] = kept
        else:
            row.pop("import_notes", None)


def _cell(record: Dict[str, Any], field: str) -> Any:
    return record.get(field)


def build_row(record: Dict[str, Any], security_type: str) -> Dict[str, Any]:
    """One mapped spreadsheet record → one `knowns.rows` entry.

    Free-text names (vesting schedule, share class, legend, document set) are
    carried through verbatim here; resolve() turns them into ids once the
    reference payload is available, or notes them and leaves them blank.
    """
    row: Dict[str, Any] = {}

    for field in ("name", "email", "notes", "custom_label"):
        text = _text(_cell(record, field))
        if text:
            row[field] = text

    relationship_raw = _cell(record, "relationship")
    if _text(relationship_raw):
        matched = _match_choice(relationship_raw, RELATIONSHIP_CHOICES)
        if matched:
            row["relationship"] = matched
        else:
            _note(row, "relationship", relationship_raw,
                  "not one of Carta's relationship options — pick one")

    kind_raw = _cell(record, "stakeholder_kind")
    if _text(kind_raw):
        kind = coerce_stakeholder_kind(kind_raw)
        if kind:
            row["stakeholder_kind"] = kind
        else:
            _note(row, "stakeholder_kind", kind_raw,
                  "not Individual or Non-individual — pick one")

    for field, label in (("quantity", "quantity"), ("issue_date", "issue date")):
        raw = _cell(record, field)
        if not _text(raw):
            continue
        value = coerce_number(raw) if field == "quantity" else coerce_date(raw)
        if value is None:
            _note(row, field, raw, "couldn't read this {} — enter it here".format(label))
        else:
            row[field] = value

    board = coerce_date(_cell(record, "board_approval_date"))
    if board:
        row["board_approval_date"] = board
    elif _text(_cell(record, "board_approval_date")):
        _note(row, "board_approval_date", _cell(record, "board_approval_date"),
              "couldn't read this board approval date — enter it here")

    for field in ("vesting_start_date",):
        raw = _cell(record, field)
        if _text(raw):
            value = coerce_date(raw)
            if value:
                row[field] = value
            else:
                _note(row, field, raw, "couldn't read this date — enter it here")

    # Names resolve() will turn into ids; staged under private keys so a row can
    # never reach the panel carrying a label where an id belongs.
    for field in ("vesting_template_id", "acceleration_template",
                  "document_set_id", "legend_id", "share_class_prefix"):
        text = _text(_cell(record, field))
        if text:
            row["_" + field] = text

    if security_type == "option_grant":
        _build_grant_fields(record, row)
    else:
        _build_cert_fields(record, row)

    return row


def _build_grant_fields(record: Dict[str, Any], row: Dict[str, Any]) -> None:
    type_raw = _cell(record, "option_type")
    if _text(type_raw):
        matched = _match_choice(type_raw, OPTION_TYPES)
        if matched:
            row["option_type"] = matched
        else:
            _note(row, "option_type", type_raw,
                  "not an option type this skill issues — pick one")

    price = _cell(record, "exercise_price")
    if _text(price):
        value = coerce_number(price)
        if value is None:
            _note(row, "exercise_price", price,
                  "couldn't read this exercise price — enter it here")
        else:
            row["exercise_price"] = value

    reason_raw = _cell(record, "grant_reason")
    if _text(reason_raw):
        matched = _match_choice(reason_raw, GRANT_REASON_CHOICES)
        if matched:
            row["grant_reason"] = matched
        else:
            _note(row, "grant_reason", reason_raw,
                  "not one of Carta's grant reasons — pick one")

    for field in ("early_exercise", "auto_exercise_at_vest", "is_hmrc_notified",
                  "is_ato_notified"):
        raw = _cell(record, field)
        if not _text(raw):
            continue
        flag = coerce_bool(raw)
        if flag is None:
            _note(row, field, raw, "couldn't read this as yes/no")
        elif flag:
            row[field] = True

    hmrc = coerce_date(_cell(record, "hmrc_notified"))
    if hmrc:
        row["hmrc_notified"] = hmrc
        row["is_hmrc_notified"] = True

    # The file's expiration wins over the plan's computed default: a shortened or
    # non-10-year term is exactly the case where the default silently disagrees
    # and still looks plausible. Left unset when the file is silent, so the
    # default applies as before.
    expiry_raw = _cell(record, "grant_expiration_date")
    if _text(expiry_raw):
        expiry = coerce_date(expiry_raw)
        if expiry is None:
            _note(row, "grant_expiration_date", expiry_raw,
                  "couldn't read this grant expiration — enter it here")
        else:
            row["grant_expiration_date"] = expiry


_CERT_ID_RE = re.compile(r"^([A-Za-z]+)[-\s]?(\d+)$")


def _build_cert_fields(record: Dict[str, Any], row: Dict[str, Any]) -> None:
    price = _cell(record, "price_per_share")
    if _text(price):
        value = coerce_number(price)
        if value is None:
            _note(row, "price_per_share", price,
                  "couldn't read this price per share — enter it here")
        else:
            row["price_per_share"] = value

    for field, label in (("cash_paid", "cash paid"),
                         ("debt_canceled", "debt canceled"),
                         ("returned_invested_capital", "returned invested capital")):
        raw = _cell(record, field)
        if not _text(raw):
            continue
        value = coerce_number(raw)
        if value is None:
            _note(row, field, raw, "couldn't read this {} — enter it here".format(label))
        else:
            row[field] = value

    r144 = _cell(record, "rule_144_date")
    if _text(r144):
        value = coerce_date(r144)
        if value is None:
            _note(row, "rule_144_date", r144,
                  "couldn't read this Rule 144 date — enter it here")
        else:
            row["rule_144_date"] = value
            issue = row.get("issue_date")
            # The panel's Rule 144 toggle is issue-date-or-other; a date that
            # differs also needs a reason, which no template column carries.
            if issue and value != issue:
                row["rule_144_mode"] = "other"
                _note(row, "rule_144_reason", value,
                      "Rule 144 date differs from the issue date — pick a reason")

    cert_id = _text(_cell(record, "prefix_number"))
    if cert_id:
        match = _CERT_ID_RE.match(cert_id)
        if match:
            row["prefix_number"] = match.group(2)
            # "CS-1" is a share-class prefix plus a number; only the number is a
            # payload field, so keep the original visible instead of silently
            # reshaping it (and let the panel show what the sheet actually said).
            _note(row, "prefix_number", cert_id,
                  "read the certificate number from {!r}".format(cert_id))
        else:
            number = coerce_number(cert_id)
            if number is None:
                _note(row, "prefix_number", cert_id,
                      "couldn't read a certificate number from this")
            else:
                row["prefix_number"] = number


# ── Parsing ──

def _records(sheet: Sheet, mapping: Dict[int, str], header_idx: int) -> List[Dict[str, Any]]:
    out = []
    for row in sheet.rows[header_idx + 1:]:
        record = {field: (row[i] if i < len(row) else None) for i, field in mapping.items()}
        if any(_text(v) for v in record.values()):
            out.append(record)
    return out


def _detect_security_type(fields: Iterable[str], sheet_name: str) -> Optional[str]:
    present = set(fields)
    grant_hits = len(present & GRANT_SIGNALS)
    cert_hits = len(present & CERT_SIGNALS)
    if grant_hits != cert_hits:
        return "option_grant" if grant_hits > cert_hits else "certificate"
    name = _norm(sheet_name)
    for hint, security_type in SHEET_NAME_HINTS:
        if _norm(hint) in name:
            return security_type
    return None


def _candidate_sheets(path: Path) -> List[Tuple[Sheet, int, Dict[int, str], List[str], str]]:
    """Every sheet with a header row, data rows, and a decidable security_type."""
    sheets = (
        [read_sheet(path)] if path.suffix.lower() in {".csv", ".tsv"}
        else _read_xlsx_sheets(path)
    )
    out = []
    for sheet in sheets:
        header_idx = find_header_row(sheet)
        if header_idx is None:
            continue
        mapping, unmapped = map_headers(sheet.rows[header_idx])
        if not _records(sheet, mapping, header_idx):
            continue
        security_type = _detect_security_type(mapping.values(), sheet.name)
        if security_type is None:
            continue
        out.append((sheet, header_idx, mapping, unmapped, security_type))
    return out


def parse_file(path: Path, sheet: Optional[str] = None) -> Parsed:
    path = Path(path)
    if not path.is_file():
        raise ParseError("No file at {}".format(path))
    suffix = path.suffix.lower()
    if suffix not in SPREADSHEET_SUFFIXES:
        raise ParseError(
            "{} isn't a spreadsheet this parser reads ({})".format(
                suffix or path.name, ", ".join(sorted(SPREADSHEET_SUFFIXES))
            )
        )

    if sheet is not None:
        target = read_sheet(path, sheet)
        header_idx = find_header_row(target)
        if header_idx is None:
            raise ParseError(
                "Couldn't find a header row in {!r} — expected Carta's importer "
                "template headers (e.g. Name, Quantity, Type).".format(target.name)
            )
        mapping, unmapped = map_headers(target.rows[header_idx])
        security_type = _detect_security_type(mapping.values(), target.name)
        if security_type is None:
            raise ParseError(
                "Couldn't tell whether {!r} holds certificates or option grants.".format(
                    target.name
                )
            )
        candidates = [(target, header_idx, mapping, unmapped, security_type)]
    else:
        candidates = _candidate_sheets(path)

    if not candidates:
        raise ParseError(
            "Nothing to import from {} — no sheet had both a recognizable header "
            "row and at least one data row.".format(path.name)
        )
    if len(candidates) > 1:
        names = [c[0].name for c in candidates]
        raise AmbiguousInput(
            "{} has more than one sheet to import: {}. Pick one.".format(
                path.name, ", ".join(names)
            ),
            names,
        )

    target, header_idx, mapping, unmapped, security_type = candidates[0]
    rows: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    plan_names: List[str] = []

    for offset, record in enumerate(_records(target, mapping, header_idx)):
        line = header_idx + 2 + offset  # 1-based sheet row
        out_of_scope = _out_of_scope(record)
        if out_of_scope:
            skipped.append({
                "row": line,
                "name": _text(_cell(record, "name")),
                "reason": "This skill issues certificates and option grants; "
                          "{} are done in the Drafts UI.".format(out_of_scope),
            })
            continue
        row = build_row(record, security_type)
        if not any(_text(v) for k, v in row.items() if k != "import_notes"):
            skipped.append({"row": line, "name": "",
                            "reason": "No usable values in this row."})
            continue
        row["row_key"] = "r{}".format(len(rows))
        rows.append(row)
        plan = _text(_cell(record, "_equity_plan"))
        if plan and plan not in plan_names:
            plan_names.append(plan)

    if not rows:
        raise ParseError(
            "Read {} row(s) from {!r} but none could be imported.".format(
                len(skipped), target.name
            )
        )

    batch_errors: List[str] = []
    if len(plan_names) > 1:
        batch_errors.append(
            "This sheet names more than one equity plan ({}). A draft set is "
            "locked to one plan — split the sheet and import each separately.".format(
                ", ".join(plan_names)
            )
        )

    return Parsed(
        security_type,
        rows,
        source_file=str(path),
        sheet=target.name,
        unmapped_columns=unmapped,
        skipped_rows=skipped,
        plan_name=plan_names[0] if plan_names else None,
        batch_errors=batch_errors,
    )


def _out_of_scope(record: Dict[str, Any]) -> Optional[str]:
    raw = _text(_cell(record, "option_type")).upper()
    if not raw:
        return None
    if _match_choice(raw, OPTION_TYPES):
        return None
    return OUT_OF_SCOPE_TYPES.get(raw)


# ── Resolution against the reference payload ──

def _unwrap(obj: Any) -> Any:
    """Peel the MCP result envelopes build_config.py's _unwrap also handles."""
    seen = 0
    while isinstance(obj, dict) and seen < 6:
        for key in ("result", "data", "content", "value"):
            if key in obj and isinstance(obj[key], (dict, list)):
                obj = obj[key]
                break
        else:
            break
        seen += 1
    return obj


def _results(raw: Any) -> List[Dict[str, Any]]:
    raw = _unwrap(raw)
    if isinstance(raw, dict):
        raw = raw.get("results", [])
    if not isinstance(raw, list):
        return []
    return [r for r in raw if isinstance(r, dict)]


def _lookup(
    records: List[Dict[str, Any]], value: str, keys: Sequence[str]
) -> Optional[Dict[str, Any]]:
    """Exact (normalized) match on any of `keys`. Deliberately not fuzzy."""
    target = _norm(value)
    if not target:
        return None
    for record in records:
        for key in keys:
            if _norm(record.get(key)) == target:
                return record
    return None


_RESOLVERS = (
    # (staged row key, reference section, match keys, id field, label)
    ("_vesting_template_id", "vesting_templates", ("name", "summary_short"), "id",
     "vesting schedule"),
    ("_acceleration_template", "acceleration_templates", ("name",), "id", "acceleration terms"),
    ("_document_set_id", "document_sets", ("name",), "id", "document set"),
    ("_legend_id", "legends", ("code", "name", "label"), "id", "legend"),
    ("_share_class_prefix", "share_classes", ("name", "prefix"), "prefix", "share class"),
)


def resolve(parsed: Parsed, reference: Dict[str, Any]) -> Parsed:
    """Turn the staged free-text names into ids, in place.

    A name that doesn't match exactly stays unset and gets an import_note, so
    the panel falls back to its own default and shows what the file said. There
    is no fuzzy fallback by design — see the module docstring.
    """
    sections = {
        name: _results(reference.get(name))
        for name in ("vesting_templates", "acceleration_templates", "document_sets",
                     "legends", "share_classes", "option_plans", "stakeholders")
    }
    roster = sections["stakeholders"]

    for row in parsed.rows:
        for staged, section, keys, id_field, label in _RESOLVERS:
            raw = row.pop(staged, None)
            if not raw:
                continue
            field = staged[1:]
            record = _lookup(sections[section], raw, keys)
            if record is None:
                _note(row, field, raw,
                      "no {} on this company matches {!r} — pick one".format(label, raw))
                continue
            value = record.get(id_field)
            if value is None:
                _note(row, field, raw,
                      "matched a {} with no {} — pick one".format(label, id_field))
                continue
            _resolved(row, field, str(value))

        name = row.get("name")
        if name:
            match = _lookup(roster, name, ("full_name", "name"))
            if match:
                # The cap-table record wins over the sheet for these three; Phase
                # 1 re-resolves authoritatively, this just pre-fills the panel.
                if match.get("email"):
                    _resolved(row, "email", match["email"])
                if match.get("kind"):
                    _resolved(row, "stakeholder_kind", match["kind"])
                if match.get("event_relationship"):
                    _resolved(row, "relationship", match["event_relationship"])

    if parsed.security_type == "option_grant" and parsed.plan_name:
        plan = _lookup(sections["option_plans"], parsed.plan_name, ("name",))
        if plan is None:
            parsed.batch_errors.append(
                "No equity plan on this company matches {!r} — pick the plan when "
                "prompted.".format(parsed.plan_name)
            )
        elif plan.get("is_expired"):
            parsed.batch_errors.append(
                "The equity plan {!r} named in this file is expired — pick another "
                "plan when prompted.".format(parsed.plan_name)
            )
        else:
            parsed.equity_plan_id = str(plan.get("id"))

    return parsed


# ── Document mode ──

def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        import fitz  # pymupdf

        with fitz.open(path) as doc:
            return "\n\n".join(
                "--- page {} ---\n{}".format(i + 1, page.get_text())
                for i, page in enumerate(doc)
            )
    if suffix == ".docx":
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for trow in table.rows:
                cells = [c.text.strip() for c in trow.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    raise ParseError("{} isn't a document this parser reads".format(suffix or path.name))


# ── Output ──

def build_import_knowns(parsed: Parsed) -> Dict[str, Any]:
    knowns: Dict[str, Any] = {
        "security_type": parsed.security_type,
        "rows": parsed.rows,
    }
    if parsed.equity_plan_id:
        knowns["equity_plan_id"] = parsed.equity_plan_id
    if parsed.batch_errors:
        knowns["batch_errors"] = list(parsed.batch_errors)
    return knowns


def build_report(parsed: Parsed) -> Dict[str, Any]:
    return {
        "mode": "spreadsheet",
        "source_file": parsed.source_file,
        "sheet": parsed.sheet,
        "security_type": parsed.security_type,
        "row_count": len(parsed.rows),
        "unmapped_columns": list(parsed.unmapped_columns),
        "skipped_rows": list(parsed.skipped_rows),
        "plan_name": parsed.plan_name,
        "batch_errors": list(parsed.batch_errors),
        "notes_by_row": {
            row["row_key"]: row.get("import_notes", [])
            for row in parsed.rows
            if row.get("import_notes")
        },
    }


def _write(out_dir: Path, name: str, payload: Any) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / name
    if isinstance(payload, str):
        path.write_text(payload, encoding="utf-8")
    else:
        path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def main(argv: Optional[List[str]] = None) -> int:
    p = argparse.ArgumentParser(
        description="Parse an uploaded spreadsheet/document into carta-issuance knowns.rows."
    )
    p.add_argument("--file", required=True, type=Path)
    p.add_argument("--sheet", help="Sheet to import when the workbook has several")
    p.add_argument("--reference", type=Path,
                   help="JSON of raw MCP reference sections (build_config's --data shape)")
    p.add_argument("--out-dir", required=True, type=Path)
    args = p.parse_args(argv)

    path: Path = args.file
    if not path.is_file():
        print("ERROR: no file at {}".format(path), file=sys.stderr)
        return 2

    suffix = path.suffix.lower()
    try:
        if suffix in DOCUMENT_SUFFIXES:
            text = extract_text(path)
            if not text.strip():
                print(
                    "ERROR: no text in {} — a scanned image needs OCR before it can "
                    "be read.".format(path.name),
                    file=sys.stderr,
                )
                return 2
            text_path = _write(args.out_dir, "_import_text.txt", text)
            report_path = _write(args.out_dir, "_import_report.json", {
                "mode": "document_text",
                "source_file": str(path),
                "characters": len(text),
                "note": "Prose has no fixed layout, so rows are not built here. "
                        "Read the text, write rows in the parser's own schema, and "
                        "mark every field confidence \"low\".",
            })
            print("IMPORT_TEXT={}".format(text_path))
            print("IMPORT_REPORT={}".format(report_path))
            return 0

        parsed = parse_file(path, args.sheet)
        if args.reference:
            parsed = resolve(parsed, json.loads(args.reference.read_text(encoding="utf-8")))
        else:
            # Drop the staged names rather than leaking a label where an id goes.
            for row in parsed.rows:
                for staged, _s, _k, _i, label in _RESOLVERS:
                    raw = row.pop(staged, None)
                    if raw:
                        _note(row, staged[1:], raw,
                              "couldn't check this {} against the company — pick one".format(label))

        knowns_path = _write(args.out_dir, "_import_knowns.json", build_import_knowns(parsed))
        report_path = _write(args.out_dir, "_import_report.json", build_report(parsed))
        print("IMPORT_KNOWNS={}".format(knowns_path))
        print("IMPORT_REPORT={}".format(report_path))
        print("ROW_COUNT={}".format(len(parsed.rows)))
        print("SECURITY_TYPE={}".format(parsed.security_type))
        return 0

    except AmbiguousInput as exc:
        print("AMBIGUOUS: {}".format(exc), file=sys.stderr)
        print("CANDIDATES={}".format(json.dumps(exc.candidates)), file=sys.stderr)
        return 2
    except ParseError as exc:
        print("ERROR: {}".format(exc), file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
